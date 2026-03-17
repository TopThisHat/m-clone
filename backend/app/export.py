"""Convert markdown research reports to DOCX format."""
from __future__ import annotations

import io
import re

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(title: str, markdown: str) -> bytes:
    """Convert a markdown string to a DOCX document and return bytes."""
    doc = Document()

    # Title
    t = doc.add_heading(title, level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            doc.add_heading(_strip_inline(text), level=min(level, 4))
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+\s*$", line):
            doc.add_paragraph("_" * 50)
            i += 1
            continue

        # Bullet list
        bullet_match = re.match(r"^[\-\*]\s+(.*)", line)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, bullet_match.group(1))
            i += 1
            continue

        # Numbered list
        num_match = re.match(r"^\d+\.\s+(.*)", line)
        if num_match:
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, num_match.group(1))
            i += 1
            continue

        # Table detection
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|", lines[i + 1]):
            table_lines = [line]
            i += 1
            # Skip separator
            if i < len(lines) and re.match(r"^\|[\s\-:|]+\|", lines[i]):
                i += 1
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            # Parse cells
            rows_data = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip().strip("|").split("|")]
                rows_data.append(cells)

            if rows_data:
                cols = max(len(r) for r in rows_data)
                table = doc.add_table(rows=len(rows_data), cols=cols)
                table.style = "Table Grid"
                for ri, row_cells in enumerate(rows_data):
                    for ci, cell_text in enumerate(row_cells):
                        if ci < cols:
                            table.rows[ri].cells[ci].text = _strip_inline(cell_text)
            continue

        # Empty line
        if line.strip() == "":
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_inline_runs(p, line)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _strip_inline(text: str) -> str:
    """Remove markdown inline formatting."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


def _add_inline_runs(paragraph, text: str) -> None:
    """Add runs to a paragraph with bold/italic support."""
    # Split on bold markers
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        bold_match = re.match(r"^\*\*(.+?)\*\*$", part)
        if bold_match:
            run = paragraph.add_run(bold_match.group(1))
            run.bold = True
        else:
            # Handle italic
            sub_parts = re.split(r"(\*.+?\*)", part)
            for sp in sub_parts:
                italic_match = re.match(r"^\*(.+?)\*$", sp)
                if italic_match:
                    run = paragraph.add_run(italic_match.group(1))
                    run.italic = True
                else:
                    # Handle links — replace [text](url) with just text
                    cleaned = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", sp)
                    if cleaned:
                        paragraph.add_run(cleaned)
